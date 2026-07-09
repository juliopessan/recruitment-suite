"""CV text extraction from uploaded files (PDF, DOCX, TXT)."""

import io
import re


class CVParseError(Exception):
    """Raised when a CV file cannot be parsed."""


def extract_cv_text(filename: str, content: bytes) -> str:
    """Extract plain text from an uploaded CV file."""
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        return _extract_pdf(content)
    if name.endswith(".docx"):
        return _extract_docx(content)
    if name.endswith((".txt", ".md")):
        return content.decode("utf-8", errors="replace")

    raise CVParseError(
        f"Unsupported file type: {filename}. Use PDF, DOCX, TXT or MD."
    )


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise CVParseError("PDF support requires the 'pypdf' package") from exc

    try:
        reader = PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise CVParseError(f"Could not read PDF: {exc}") from exc

    if not text.strip():
        raise CVParseError("PDF contains no extractable text (scanned image?)")
    return text


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise CVParseError("DOCX support requires the 'python-docx' package") from exc

    try:
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(p for p in parts if p and p.strip())
    except Exception as exc:
        raise CVParseError(f"Could not read DOCX: {exc}") from exc


EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
YEARS_RE = re.compile(r"(\d{1,2})\s*\+?\s*(?:years?|anos?)", re.IGNORECASE)


def guess_candidate_fields(cv_text: str) -> dict:
    """Best-effort extraction of name/email/experience from CV text."""
    fields: dict = {}

    email = EMAIL_RE.search(cv_text)
    if email:
        fields["email"] = email.group(0)

    years = [int(m) for m in YEARS_RE.findall(cv_text) if int(m) <= 50]
    if years:
        fields["total_years_experience"] = max(years)

    # First reasonable-looking line as the name
    for line in cv_text.splitlines():
        line = line.strip()
        if 2 <= len(line.split()) <= 5 and len(line) < 60 and not EMAIL_RE.search(line) \
                and not any(ch.isdigit() for ch in line):
            fields["name"] = line
            break

    return fields
